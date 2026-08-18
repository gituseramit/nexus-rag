from pypdf import PdfReader
from docx import Document
import csv

class TextExtractor:
    def extract(self, file_path: str, mime_type: str) -> tuple[str, dict]:
        metadata = {}
        text = ""
        
        if file_path.endswith('.pdf') or mime_type == 'application/pdf':
            reader = PdfReader(file_path)
            metadata["page_count"] = len(reader.pages)
            if reader.metadata:
                metadata["author"] = reader.metadata.author
                metadata["title"] = reader.metadata.title
            text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            
        elif file_path.endswith('.docx') or mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs]
            tables = []
            for table in doc.tables:
                for row in table.rows:
                    tables.append(" | ".join([cell.text for cell in row.cells]))
            text = "\n".join(paragraphs + tables)
            
        elif file_path.endswith('.csv') or mime_type == 'text/csv':
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                text = "\n".join([" | ".join(row) for row in reader])
                
        elif file_path.endswith('.txt') or file_path.endswith('.md') or 'text' in str(mime_type):
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                
        else:
            raise ValueError(f"Unsupported file type: {mime_type} or {file_path}")
            
        # Clean text
        text = text.replace('\x00', '')
        text = "\n".join([line for line in text.splitlines() if line.strip()])
        
        return text, metadata

extractor = TextExtractor()
